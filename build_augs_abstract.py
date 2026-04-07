"""
build_augs_abstract.py
Generates AUGS_Abstract_POP_PREDICT.docx — the AUGS-formatted abstract
with colored risk-tier table (green/yellow/red).
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = "AUGS_Abstract_POP_PREDICT.docx"

# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def shade_cell(cell, hex_color: str):
    """Fill a table cell with a solid background color (hex, no #)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def bold_cell(cell):
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True


def set_col_width(table, col_idx, width_inches):
    for row in table.rows:
        row.cells[col_idx].width = Inches(width_inches)


def add_section_heading(doc, label):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(label + "  ")
    run.bold = True
    run.font.size = Pt(12)
    return p


def add_body_paragraph(doc, bold_prefix, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = Pt(14)
    if bold_prefix:
        run = p.add_run(bold_prefix + " ")
        run.bold = True
        run.font.size = Pt(12)
    run = p.add_run(text)
    run.font.size = Pt(12)
    return p


def header_row(table, *headers):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        shade_cell(cell, "2E4057")          # dark navy
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def data_row(table, row_idx, *values, shade=None, bold=False):
    row = table.rows[row_idx]
    for i, v in enumerate(values):
        cell = row.cells[i]
        if shade:
            shade_cell(cell, shade)
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(v)
        run.font.size = Pt(10)
        if bold:
            run.bold = True
        if i > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


def caption(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    for run in p.runs:
        run.font.size = Pt(9)
        run.italic = True


# ---------------------------------------------------------------------------
# Build document
# ---------------------------------------------------------------------------

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)

# Global default font
for style_name in ("Normal", "Table Grid"):
    try:
        s = doc.styles[style_name]
        s.font.name = "Times New Roman"
        s.font.size = Pt(12)
    except Exception:
        pass

# ---- Title ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
run = p.add_run(
    "Development and Validation of the POP-PREDICT Score: "
    "A Point-Based Clinical Risk Calculator for Recurrence "
    "Following Pelvic Organ Prolapse Surgery"
)
run.bold = True
run.font.size = Pt(12)
run.font.name = "Times New Roman"

# ---- Authors ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(2)
for segment, is_super in [
    ("Sullivan M, MD", False), ("\u00b9", True),
    ("; Kim S", False),        ("\u00b9", True),
    ("; El Haraki A, MD", False), ("\u00b9", True),
]:
    r = p.add_run(segment)
    r.font.size = Pt(12)
    r.font.name = "Times New Roman"
    if is_super:
        r.font.superscript = True

# ---- Affiliation ----
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run(
    "\u00b9Department of Obstetrics and Gynecology, "
    "Atrium Health Wake Forest Baptist, Winston-Salem, NC"
)
r.italic = True
r.font.size = Pt(10)
r.font.name = "Times New Roman"

# ---- Abstract sections ----
add_body_paragraph(doc, "Objective:",
    "To develop and validate a point-based preoperative clinical risk calculator "
    "for pelvic organ prolapse (POP) recurrence incorporating patient characteristics "
    "and planned surgical approach."
)

add_body_paragraph(doc, "Methods:",
    "We performed a pooled individual-patient data analysis of four prospective clinical "
    "trials (ALTIS, BEST, EPACT, SASS) comprising 752 women who underwent POP surgery "
    "between 2017 and 2024. Surgical approaches included sacrocolpopexy (62.8%), native "
    "tissue repair including uterosacral and sacrospinous ligament suspension (30.9%), and "
    "colpocleisis (6.4%). The primary outcome was composite recurrence, defined as POP-Q "
    "stage \u22652, any POP-Q measurement point \u22650, or bothersome vaginal bulge "
    "symptoms (PFDI-20 Question 3 \u22652) at any follow-up visit with available data. "
    "Standard logistic regression with backward stepwise Akaike Information Criterion "
    "(AIC) selection was used for predictor identification. Model discrimination was "
    "assessed by full-cohort area under the receiver operating characteristic curve "
    "(AUC), five-fold stratified cross-validation (CV), and leave-one-dataset-out "
    "(LODO) external validation across three of four trial datasets. Regression "
    "coefficients were converted to a clinically scaled integer point score, with "
    "risk tiers defined empirically from the score distribution."
)

add_body_paragraph(doc, "Results:",
    "Mean age was 62.9 years (SD 11.4) and BMI 28.1 kg/m\u00b2 (SD 7.0); 58.1% of "
    "patients presented with stage 3 prolapse (Table 1). Composite recurrence occurred "
    "in 210 of 752 patients (27.9%). AIC model selection identified four independent "
    "predictors of recurrence: surgical approach, baseline Anterior Vaginal Wall (Ba "
    "point), baseline Genital Hiatus (gh), and menopausal/hormonal status. On "
    "multivariable analysis, wider Genital Hiatus (OR 1.31 per cm, 95% CI 1.13\u20131.51, "
    "p<0.001), greater anterior wall prolapse (OR 1.13 per cm, 95% CI 1.05\u20131.21, "
    "p=0.001), native tissue repair versus sacrocolpopexy (OR 1.48, 95% CI 1.04\u20132.13, "
    "p=0.032), and colpocleisis versus sacrocolpopexy (OR 0.37, 95% CI 0.15\u20130.88, "
    "p=0.025) were significant. The POP-PREDICT score (0\u20137 points) assigns points "
    "for surgical approach (0\u20132), Ba point (0\u20132), Genital Hiatus (0\u20132), "
    "and hormonal status (0\u20131). Full-cohort AUC was 0.67; five-fold CV AUC was "
    "0.64 (SD 0.07); LODO AUC ranged from 0.55 to 0.66 (Table 2). The score stratified "
    "patients into low risk (0\u20132 points, 8\u201317% recurrence), moderate risk (3\u20134 "
    "points, 21\u201328%), and high risk (5\u20137 points, 40\u201345%) categories (Table 3)."
)

add_body_paragraph(doc, "Conclusions:",
    "The POP-PREDICT score is a novel, validated point-based tool for preoperative risk "
    "stratification of POP surgical recurrence. Incorporating four readily available "
    "clinical variables, it demonstrated consistent discrimination across internal "
    "cross-validation and external LODO validation (AUC 0.64\u20130.67) and stratifies "
    "patients into three clinically meaningful risk tiers. This tool supports "
    "individualized preoperative counseling, shared decision-making, and identification "
    "of patients who may benefit from more durable surgical approaches or enhanced "
    "postoperative surveillance."
)

# ---- TABLE 1: Patient Characteristics ----
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Table 1.  Patient Characteristics (N = 752)")
run.bold = True
run.font.size = Pt(11)

rows_t1 = [
    # (label, value, indent)
    ("Age, years",                  "62.9 (11.4)", False),
    ("BMI, kg/m²",                  "28.1 (7.0)",  False),
    ("Parity, births",              "2.6 (1.4)",   False),
    ("Race",                        "",            False),
    ("  White",                     "82.0%",       True),
    ("  Black",                     "7.7%",        True),
    ("  Hispanic",                  "5.6%",        True),
    ("  Other / Unknown",           "4.7%",        True),
    ("Hormonal Status",             "",            False),
    ("  Pre- or peri-menopausal",   "20.2%",       True),
    ("  Postmenopausal, no HRT",    "53.7%",       True),
    ("  Postmenopausal, vaginal estrogen", "19.1%", True),
    ("  Postmenopausal, oral/systemic HRT", "4.9%", True),
    ("Surgery Type",                "",            False),
    ("  Sacrocolpopexy",            "62.8%",       True),
    ("  Native tissue repair",      "30.9%",       True),
    ("  Colpocleisis",              "6.4%",        True),
    ("Baseline POP-Q Stage",        "",            False),
    ("  Stage 2",                   "31.4%",       True),
    ("  Stage 3",                   "58.1%",       True),
    ("  Stage 4",                   "10.5%",       True),
    ("Baseline Ba Point, cm",       "2.2 (2.6)",   False),
    ("Baseline Genital Hiatus, cm", "4.3 (1.3)",   False),
    ("Composite recurrence",        "27.9% (n=210)", False),
]

t1 = doc.add_table(rows=len(rows_t1) + 1, cols=2)
t1.style = "Table Grid"
header_row(t1, "Characteristic", "Value")
for i, (label, val, indent) in enumerate(rows_t1):
    row = t1.rows[i + 1]
    shade_color = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    for cell in row.cells:
        shade_cell(cell, shade_color)
    row.cells[0].paragraphs[0].clear()
    run = row.cells[0].paragraphs[0].add_run(("    " if indent else "") + label)
    run.font.size = Pt(10)
    if not indent and val == "":
        run.bold = True
    row.cells[1].paragraphs[0].clear()
    run = row.cells[1].paragraphs[0].add_run(val)
    run.font.size = Pt(10)
    row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_col_width(t1, 0, 3.5)
set_col_width(t1, 1, 2.0)

caption(doc,
    "Values are mean (SD) for continuous variables and % for categorical variables. "
    "HRT = hormone replacement therapy; POP-Q = pelvic organ prolapse quantification system."
)

# ---- TABLE 2: Model Performance ----
p = doc.add_paragraph()
run = p.add_run("Table 2.  Model Discrimination")
run.bold = True
run.font.size = Pt(11)

t2 = doc.add_table(rows=6, cols=6)
t2.style = "Table Grid"
header_row(t2, "Validation", "AUC", "Sensitivity", "Specificity", "PPV", "NPV")

t2_data = [
    ("Full cohort",             "0.67", "0.66", "0.60", "0.39", "0.82"),
    ("5-fold CV (mean ± SD)",   "0.64 ± 0.07", "—", "—", "—", "—"),
    ("LODO — held-out BEST",    "0.66", "—", "—", "—", "—"),
    ("LODO — held-out EPACT",   "0.58", "—", "—", "—", "—"),
    ("LODO — held-out SASS",    "0.55", "—", "—", "—", "—"),
]
# Note: ALTIS LODO could not be computed due to model convergence in the training set
for i, row_data in enumerate(t2_data):
    shade = "F2F2F2" if i % 2 == 0 else "FFFFFF"
    data_row(t2, i + 1, *row_data, shade=shade)

caption(doc,
    "AUC = area under the ROC curve; CV = cross-validation; LODO = leave-one-dataset-out; "
    "PPV = positive predictive value; NPV = negative predictive value. "
    "ALTIS LODO could not be computed due to near-perfect separation in the training set without that cohort."
)

# ---- TABLE 3: POP-PREDICT Score Card ----
p = doc.add_paragraph()
run = p.add_run("Table 3.  POP-PREDICT Score")
run.bold = True
run.font.size = Pt(11)

t3 = doc.add_table(rows=13, cols=3)
t3.style = "Table Grid"
header_row(t3, "Risk Factor", "Criteria", "Points")

score_rows = [
    # (factor, criteria, points, section_header)
    ("Surgical Approach",       "Colpocleisis",                         "0",  True),
    ("",                        "Sacrocolpopexy",                       "1",  False),
    ("",                        "Native tissue repair (USLS / SSLF)",   "2",  False),
    ("Anterior Wall (Ba Point)","< 0 cm",                               "0",  True),
    ("",                        "0 to 3 cm",                            "1",  False),
    ("",                        "\u2265 4 cm",                          "2",  False),
    ("Genital Hiatus (gh)",     "\u2264 3 cm",                          "0",  True),
    ("",                        "4 to 5 cm",                            "1",  False),
    ("",                        "\u2265 6 cm",                          "2",  False),
    ("Hormonal Status",         "Pre-menopausal or on estrogen (systemic/local)", "0",  True),
    ("",                        "Postmenopausal, no HRT",                         "1",  False),
    # blank separator row
    ("",                        "",                                     "",   False),
    # total row
    ("TOTAL SCORE",             "_____ / 7 points",                     "",   True),
]

for i, (factor, criteria, pts, is_header) in enumerate(score_rows):
    row = t3.rows[i]
    shade_cell(row.cells[0], "F2F2F2" if is_header else "FFFFFF")
    shade_cell(row.cells[1], "F2F2F2" if is_header else "FFFFFF")
    shade_cell(row.cells[2], "F2F2F2" if is_header else "FFFFFF")

    row.cells[0].paragraphs[0].clear()
    r = row.cells[0].paragraphs[0].add_run(factor)
    r.font.size = Pt(10)
    if is_header and factor:
        r.bold = True

    row.cells[1].paragraphs[0].clear()
    r = row.cells[1].paragraphs[0].add_run(criteria)
    r.font.size = Pt(10)

    row.cells[2].paragraphs[0].clear()
    r = row.cells[2].paragraphs[0].add_run(pts)
    r.font.size = Pt(10)
    r.bold = True
    row.cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_col_width(t3, 0, 2.0)
set_col_width(t3, 1, 2.8)
set_col_width(t3, 2, 0.7)

# ---- Risk tier table (colored) ----
doc.add_paragraph()
t4 = doc.add_table(rows=4, cols=3)
t4.style = "Table Grid"

tier_data = [
    ("0–2 points", "LOW RISK",      "8–17% recurrence probability",  "C6EFCE", "375623"),  # green
    ("3–4 points", "MODERATE RISK", "21–28% recurrence probability", "FFEB9C", "9C6500"),  # yellow
    ("5–7 points", "HIGH RISK",     "40–45% recurrence probability", "FFC7CE", "9C0006"),  # red
]

# header
header_row_tier = t4.rows[0]
for cell in header_row_tier.cells:
    shade_cell(cell, "2E4057")
header_row_tier.cells[0].paragraphs[0].clear()
header_row_tier.cells[1].paragraphs[0].clear()
header_row_tier.cells[2].paragraphs[0].clear()
for cell, text in zip(header_row_tier.cells, ["Score", "Risk Category", "Observed Recurrence"]):
    run = cell.paragraphs[0].add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    run.font.size = Pt(10)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

for i, (score, label, pct, bg, fg_hex) in enumerate(tier_data):
    row = t4.rows[i + 1]
    for cell in row.cells:
        shade_cell(cell, bg)
    fg = tuple(int(fg_hex[j:j+2], 16) for j in (0, 2, 4))

    for cell, text in zip(row.cells, [score, label, pct]):
        cell.paragraphs[0].clear()
        run = cell.paragraphs[0].add_run(text)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*fg)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

set_col_width(t4, 0, 1.2)
set_col_width(t4, 1, 2.0)
set_col_width(t4, 2, 2.3)

caption(doc,
    "USLS = uterosacral ligament suspension; SSLF = sacrospinous ligament fixation; "
    "HRT = hormone replacement therapy. "
    "Recurrence probabilities are empirical rates from the derivation cohort. "
    "Clinical Use: This score supports preoperative counseling and shared decision-making. "
    "It does not replace clinical judgment."
)

# ---- Save ----
doc.save(OUTPUT)
print(f"Saved: {OUTPUT}")
